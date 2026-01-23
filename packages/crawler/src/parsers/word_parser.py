"""Word文档解析器"""

import os
import re
from typing import Any, Dict, List, Optional

from loguru import logger

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    Document = None
    Table = None


class WordParser:
    """Word文档解析器

    功能：
    - 使用python-docx读取Word文档
    - 表格提取
    - 段落提取
    - 图片提取（需OCR）
    """

    # 标准字段映射
    FIELD_MAPPING = {
        "招录机关": "department_name",
        "部门名称": "department_name",
        "职位名称": "position_name",
        "岗位名称": "position_name",
        "职位代码": "position_code",
        "招录人数": "recruit_count",
        "学历": "education_min",
        "专业": "major_specific",
        "政治面貌": "political_status",
        "年龄": "age_requirement",
        "工作经历": "work_exp_requirement",
        "户籍": "hukou_requirement",
        "性别": "gender_required",
        "工作地点": "work_location",
        "备注": "notes",
    }

    def __init__(self, ocr_parser=None):
        """
        Args:
            ocr_parser: OCR解析器实例（用于图片识别）
        """
        self.ocr_parser = ocr_parser

        if Document is None:
            logger.warning("python-docx未安装，部分功能将不可用")

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """解析Word文档

        Args:
            file_path: 文件路径

        Returns:
            解析后的数据列表
        """
        if Document is None:
            raise ImportError("需要安装python-docx库")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 检查文件类型
        if file_path.lower().endswith(".doc"):
            # 旧版doc格式，需要转换
            file_path = self._convert_doc_to_docx(file_path)
            if not file_path:
                raise ValueError("无法处理.doc格式文件，请转换为.docx格式")

        doc = Document(file_path)

        all_results = []

        # 解析所有表格
        for table_idx, table in enumerate(doc.tables):
            if self._is_position_table(table):
                results = self._parse_table(table)
                logger.info(f"表格{table_idx+1}: {len(results)}条记录")
                all_results.extend(results)

        return all_results

    def _convert_doc_to_docx(self, file_path: str) -> Optional[str]:
        """转换doc到docx格式"""
        try:
            import subprocess

            output_path = file_path + "x"

            # 尝试使用LibreOffice转换
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    os.path.dirname(file_path),
                    file_path,
                ],
                capture_output=True,
                timeout=30,
            )

            if os.path.exists(output_path):
                return output_path

        except Exception as e:
            logger.error(f"doc转换失败: {e}")

        return None

    def _is_position_table(self, table) -> bool:
        """判断是否为职位表"""
        if not table.rows:
            return False

        # 获取第一行文本
        first_row = table.rows[0]
        header_text = " ".join(cell.text for cell in first_row.cells)

        position_keywords = ["职位", "岗位", "部门", "学历", "专业", "人数"]
        match_count = sum(1 for kw in position_keywords if kw in header_text)

        return match_count >= 3

    def _parse_table(self, table) -> List[Dict[str, Any]]:
        """解析表格"""
        if not table.rows:
            return []

        # 获取表头
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]

        # 去重（合并单元格可能导致重复）
        headers = self._deduplicate_headers(headers)

        # 映射字段名
        mapped_headers = self._map_headers(headers)

        # 解析数据行
        results = []

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]

            # 去重
            cells = self._deduplicate_cells(cells, headers)

            if not any(cells):
                continue

            item = {}
            for col_idx, cell in enumerate(cells):
                if col_idx < len(mapped_headers):
                    field_name = mapped_headers[col_idx]
                    if field_name and cell:
                        item[field_name] = cell

            if item:
                results.append(item)

        return results

    def _deduplicate_headers(self, headers: List[str]) -> List[str]:
        """去除重复的表头（处理合并单元格）"""
        result = []
        prev = None

        for header in headers:
            if header != prev:
                result.append(header)
                prev = header

        return result

    def _deduplicate_cells(self, cells: List[str], headers: List[str]) -> List[str]:
        """去除重复的单元格值"""
        if len(cells) == len(headers):
            return cells

        # 如果单元格数量超过表头，可能是合并单元格
        result = []
        prev = None

        for cell in cells:
            if cell != prev:
                result.append(cell)
                prev = cell

        # 补齐长度
        while len(result) < len(headers):
            result.append("")

        return result[: len(headers)]

    def _map_headers(self, headers: List[str]) -> List[Optional[str]]:
        """映射表头到标准字段名"""
        mapped = []

        for header in headers:
            if not header:
                mapped.append(None)
                continue

            # 精确匹配
            if header in self.FIELD_MAPPING:
                mapped.append(self.FIELD_MAPPING[header])
                continue

            # 模糊匹配
            matched = False
            for key, value in self.FIELD_MAPPING.items():
                if key in header or header in key:
                    mapped.append(value)
                    matched = True
                    break

            if not matched:
                field_name = re.sub(r"\s+", "_", header.lower())
                mapped.append(field_name)

        return mapped

    def extract_text(self, file_path: str) -> str:
        """提取文档全部文本"""
        if Document is None:
            raise ImportError("需要安装python-docx库")

        doc = Document(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        return "\n\n".join(paragraphs)

    def extract_images(self, file_path: str, output_dir: str) -> List[str]:
        """提取文档中的图片"""
        if Document is None:
            raise ImportError("需要安装python-docx库")

        doc = Document(file_path)
        image_paths = []

        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                image_data = rel.target_part.blob

                # 确定图片格式
                content_type = rel.target_part.content_type
                ext = content_type.split("/")[-1]

                # 保存图片
                img_name = f"image_{len(image_paths)+1}.{ext}"
                img_path = os.path.join(output_dir, img_name)

                with open(img_path, "wb") as f:
                    f.write(image_data)

                image_paths.append(img_path)

        return image_paths
